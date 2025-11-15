"""
Utilitaires pour générer des documents Word (.docx) à partir de Markdown.
"""
from docx import Document
from docxcompose.composer import Composer
from markdown_it import MarkdownIt
from django.conf import settings
from pathlib import Path
import tempfile
import re
from docx.shared import Pt


COVER_PATH = getattr(settings, "KIT_COVER_PATH", "assets/Kit_Complet_Preparation_Couverture_BSG.docx")


def markdown_to_docx_body(md_text: str) -> Document:
    """
    Convertit un Markdown simple en Paragraphs/Tables rudimentaires.
    NOTE: Implémentation minimale (titres, paragraphes, listes, tableaux simples).
    Pour un rendu plus riche, étendre ce parseur.
    """
    md = MarkdownIt()
    tokens = md.parse(md_text or "")
    doc = Document()
    buf_list = []
    current_list_items = []

    i = 0
    while i < len(tokens):
        t = tokens[i]
        
        if t.type == "heading_open":
            level = int(t.tag[-1]) if t.tag.startswith('h') else 1
            # Flush buffer si besoin
            if buf_list:
                p = doc.add_paragraph("\n".join(buf_list))
                buf_list = []
            
            # Chercher le contenu du titre dans les tokens suivants
            i += 1
            title_text = []
            while i < len(tokens) and tokens[i].type != "heading_close":
                if tokens[i].type == "inline":
                    title_text.append(tokens[i].content.strip())
                i += 1
            
            if title_text:
                p = doc.add_paragraph()
                run = p.add_run(" ".join(title_text))
                run.bold = True
                # Avoid arithmetic on style sizes (can be None or non-scalar).
                # Use explicit point sizes for headings.
                if level == 1:
                    run.font.size = Pt(18)
                elif level == 2:
                    run.font.size = Pt(14)
        
        elif t.type == "table_open":
            # Tableau Markdown
            table_lines = []
            i += 1
            while i < len(tokens) and tokens[i].type != "table_close":
                if tokens[i].type == "tr":
                    row_tokens = []
                    j = i + 1
                    while j < len(tokens) and tokens[j].type != "tr_close":
                        if tokens[j].type == "inline":
                            row_tokens.append(tokens[j].content.strip())
                        j += 1
                    if row_tokens:
                        table_lines.append(row_tokens)
                    i = j
                i += 1
            
            if table_lines and len(table_lines) > 1:
                # Première ligne = header
                table = doc.add_table(rows=len(table_lines) - 1, cols=len(table_lines[0]))
                # Headers
                for col_idx, header_text in enumerate(table_lines[0]):
                    table.rows[0].cells[col_idx].text = header_text
                    table.rows[0].cells[col_idx].paragraphs[0].runs[0].bold = True
                
                # Data rows
                for row_idx, row_data in enumerate(table_lines[1:], start=1):
                    for col_idx, cell_text in enumerate(row_data):
                        if row_idx < len(table.rows):
                            table.rows[row_idx].cells[col_idx].text = cell_text
            continue
        
        elif t.type == "bullet_list_open":
            current_list_items = []
        
        elif t.type == "list_item_open":
            item_text = []
            i += 1
            while i < len(tokens) and tokens[i].type != "list_item_close":
                if tokens[i].type == "inline":
                    item_text.append(tokens[i].content.strip())
                i += 1
            if item_text:
                current_list_items.append(" ".join(item_text))
        
        elif t.type == "bullet_list_close":
            for item in current_list_items:
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(item)
            current_list_items = []
        
        elif t.type == "paragraph_open":
            # Commencer un paragraphe
            buf_list = []
        
        elif t.type == "inline":
            text = t.content.strip()
            if text:
                buf_list.append(text)
        
        elif t.type == "paragraph_close":
            if buf_list:
                doc.add_paragraph("\n".join(buf_list))
                buf_list = []
        
        elif t.type == "blockquote_open":
            # Citation
            quote_lines = []
            i += 1
            while i < len(tokens) and tokens[i].type != "blockquote_close":
                if tokens[i].type == "paragraph_open":
                    i += 1
                    while i < len(tokens) and tokens[i].type != "paragraph_close":
                        if tokens[i].type == "inline":
                            quote_lines.append(tokens[i].content.strip())
                        i += 1
                i += 1
            
            if quote_lines:
                p = doc.add_paragraph()
                p.style = "Quote"
                p.add_run(" ".join(quote_lines))
            continue
        
        i += 1

    if buf_list:
        doc.add_paragraph("\n".join(buf_list))

    return doc


def build_docx_with_cover(md_text: str, out_dir: Path | None = None) -> str:
    """
    Construit un document Word avec couverture + contenu Markdown.
    
    Args:
        md_text: Texte Markdown à convertir
        out_dir: Répertoire de sortie (défaut: répertoire temporaire)
    
    Returns:
        Chemin vers le fichier DOCX final
    """
    body = markdown_to_docx_body(md_text)
    if out_dir is None:
        out_dir = Path(tempfile.gettempdir())
    else:
        out_dir = Path(out_dir)
        out_dir.mkdir(parents=True, exist_ok=True)

    body_path = out_dir / "kit_body.docx"
    body.save(str(body_path))

    cover_path = Path(COVER_PATH)
    if not cover_path.exists():
        # Pas de couverture -> retourner le body
        return str(body_path)

    cover_doc = Document(str(cover_path))
    composer = Composer(cover_doc)
    body_doc = Document(str(body_path))
    composer.append(body_doc)

    final_path = out_dir / "kit_final.docx"
    composer.save(str(final_path))
    return str(final_path)

