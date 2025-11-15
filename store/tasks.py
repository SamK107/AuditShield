"""
Tâches Celery pour le traitement des kits
"""
import logging
import unicodedata
from pathlib import Path
from django.conf import settings
from django.template.loader import render_to_string
from django.utils import timezone
from celery import shared_task
from openai import OpenAI
from store.models import ClientInquiry, InquiryDocument

logger = logging.getLogger(__name__)


def _normalize_ascii(value):
    """
    Normalise une valeur pour qu'elle soit ASCII-safe.
    Remplace les caractères Unicode problématiques par leurs équivalents ASCII.
    Utilisé pour les paramètres optionnels du client OpenAI (ORG, PROJECT)
    qui pourraient contenir des apostrophes courbes ou autres caractères
    Unicode.
    """
    if not value:
        return value
    if isinstance(value, str):
        # Normaliser les apostrophes courbes (') et autres caractères Unicode
        # en leurs équivalents ASCII (')
        # Exemple: ' → '
        value = value.replace('\u2019', "'")  # Apostrophe courbe droite
        value = value.replace('\u2018', "'")  # Apostrophe courbe gauche
        value = value.replace('\u201C', '"')  # Guillemet gauche
        value = value.replace('\u201D', '"')  # Guillemet droit
        # Normalisation Unicode générale
        normalized = unicodedata.normalize('NFKD', value)
        # Encoder en ASCII en ignorant les caractères non-ASCII restants
        ascii_safe = normalized.encode('ascii', 'ignore').decode('ascii')
        return ascii_safe
    return value


@shared_task(bind=True, max_retries=3)
def build_kit_word(self, inquiry_id):
    """
    Tâche Celery pour générer le document Word du kit via OpenAI + docxtpl.
    Ne démarre QUE si le paiement est confirmé (PAID).
    
    Args:
        inquiry_id: ID de l'instance ClientInquiry
    """
    from django.db import transaction
    from store.models import GeneratedDraft
    from docxtpl import DocxTemplate
    from django.core.files.base import ContentFile
    
    with transaction.atomic():
        inquiry = ClientInquiry.objects.select_for_update().get(id=inquiry_id)
        
        # Guard : vérifier le paiement
        if inquiry.payment_status != "PAID":
            logger.warning(f"[build_kit_word] Paiement non confirmé pour inquiry {inquiry_id}")
            return
        
        # Guard : vérifier l'état de traitement
        if inquiry.processing_state not in ("PAID", "IA_RUNNING"):
            logger.warning(f"[build_kit_word] État invalide: {inquiry.processing_state} pour inquiry {inquiry_id}")
            return
        
        # Mettre à jour l'état
        inquiry.processing_state = "IA_RUNNING"
        inquiry.save(update_fields=["processing_state"])
    
    try:
        # Lire tous les fichiers uploadés
        documents = InquiryDocument.objects.filter(inquiry=inquiry)
        file_contents = []
        for doc in documents:
            try:
                with doc.file.open('rb') as f:
                    content = f.read()
                    file_contents.append({
                        'name': doc.original_name or doc.file.name,
                        'content': content,
                        'size': len(content)
                    })
            except Exception as e:
                logger.warning(f"Erreur lecture fichier {doc.id}: {e}")
        
        # Construire le prompt avec les données du formulaire
        prompt_md = render_to_string(
            "ai/prompts/kit_complet_consigne.md",
            {"inquiry": inquiry, "documents": file_contents}
        )
        
        # Vérifier que la clé API est définie
        if not settings.OPENAI_API_KEY:
            error_msg = "OPENAI_API_KEY n'est pas configurée dans les settings"
            logger.error(error_msg)
            raise ValueError(error_msg)
        
        # Configuration du client OpenAI avec options optionnelles
        # Ne pas normaliser l'API key (doit rester exacte)
        # Mais normaliser les paramètres optionnels qui pourraient
        # contenir des caractères Unicode problématiques
        client_kwargs = {"api_key": settings.OPENAI_API_KEY}

        # Ajouter OPENAI_ORG si défini (déjà nettoyé dans settings)
        if hasattr(settings, 'OPENAI_ORG') and settings.OPENAI_ORG:
            org_normalized = _normalize_ascii(settings.OPENAI_ORG)
            if org_normalized and org_normalized.strip():
                client_kwargs["organization"] = org_normalized

        # Ajouter OPENAI_PROJECT si défini (déjà nettoyé dans settings)
        if hasattr(settings, 'OPENAI_PROJECT') and settings.OPENAI_PROJECT:
            project_normalized = _normalize_ascii(settings.OPENAI_PROJECT)
            if project_normalized and project_normalized.strip():
                client_kwargs["project"] = project_normalized
        if hasattr(settings, 'OPENAI_BASE_URL') and settings.OPENAI_BASE_URL:
            # Les URLs doivent rester exactes (pas de normalisation)
            client_kwargs["base_url"] = settings.OPENAI_BASE_URL

        # Appeler OpenAI avec le modèle configuré
        client = OpenAI(**client_kwargs)
        
        # Préparer les messages pour OpenAI
        # S'assurer que le prompt est en UTF-8 valide
        if isinstance(prompt_md, bytes):
            prompt_md = prompt_md.decode('utf-8', errors='replace')
        
        messages = [
            {
                "role": "system",
                "content": (
                    "Tu es un expert en audit et conformite pour les "
                    "administrations publiques. Tu generes des documents "
                    "structures et professionnels en francais."
                )
            },
            {
                "role": "user",
                "content": prompt_md
            }
        ]
        
        # Ajouter le contenu des fichiers si disponibles
        if file_contents:
            for file_info in file_contents[:10]:  # Max 10 fichiers
                try:
                    # Extraire le texte des fichiers (simplifié)
                    if file_info['name'].endswith('.txt'):
                        text_content = file_info['content'].decode(
                            'utf-8', errors='ignore'
                        )
                        messages.append({
                            "role": "user",
                            "content": (
                                f"Contenu du fichier {file_info['name']}:"
                                f"\n\n{text_content}"
                            )
                        })
                except Exception as e:
                    logger.warning(
                        f"Erreur traitement fichier {file_info['name']}: {e}"
                    )
        
        # Déterminer le modèle à utiliser (avec fallbacks)
        model = getattr(settings, 'OPENAI_CHAT_MODEL', 'gpt-4o-mini')
        fallbacks = getattr(
            settings,
            'OPENAI_CHAT_MODEL_FALLBACKS',
            ['gpt-4o', 'gpt-4o-mini']
        )
        models_to_try = [model] + [f for f in fallbacks if f != model]

        # Appel API OpenAI avec gestion des fallbacks
        response = None
        last_error = None
        for model_name in models_to_try:
            try:
                response = client.chat.completions.create(
                    model=model_name,
                    messages=messages,
                    temperature=0.7,
                    max_tokens=8000
                )
                logger.info(
                    f"Modèle OpenAI utilisé avec succès: {model_name}"
                )
                break
            except Exception as e:
                last_error = e
                logger.warning(f"Erreur avec le modèle {model_name}: {e}")
                if model_name == models_to_try[-1]:
                    # Dernier modèle, on propage l'erreur
                    raise

        if response is None:
            raise last_error or ValueError(
                "Aucun modèle OpenAI n'a fonctionné"
            )
        
        generated_md = response.choices[0].message.content
        
        # 1) Agréger les données (payload + InquiryDocument)
        context_data = {
            "inquiry": inquiry,
            "organization_name": inquiry.organization_name or "",
            "contact_name": inquiry.contact_name or "",
            "email": inquiry.email,
            "statut_juridique": inquiry.statut_juridique or "",
            "location": inquiry.location or "",
            "sector": inquiry.sector or "",
            "mission_text": inquiry.mission_text or "",
            "context_text": inquiry.context_text or "",
            "budget_range": inquiry.budget_range or "",
            "funding_sources": ", ".join(inquiry.funding_sources or []),
            "audits_types": ", ".join(inquiry.audits_types or []),
            "audits_frequency": inquiry.audits_frequency or "",
            "staff_size": inquiry.staff_size or "",
            "org_chart_text": inquiry.org_chart_text or "",
            "notes_text": inquiry.notes_text or "",
            "generated_content": generated_md,
            "documents_count": len(file_contents),
        }
        
        # 2) docxtpl: charger template, rendu, sauvegarder fichier
        # TODO: Récupérer le template depuis la config ou un chemin par défaut
        template_path = Path(settings.PRIVATE_MEDIA_ROOT) / "templates" / "kit_complet_template.docx"
        if not template_path.exists():
            # Fallback: créer un template minimal ou utiliser un template par défaut
            logger.warning(f"Template introuvable: {template_path}, utilisation d'un template par défaut")
            # Créer un template minimal avec docx
            from docx import Document
            doc = Document()
            doc.add_heading("Kit Complet de Préparation à l'Audit", 0)
            doc.add_paragraph(f"Organisation: {context_data['organization_name']}")
            doc.add_paragraph(f"Contenu généré:\n{generated_md}")
            temp_path = Path(settings.PRIVATE_MEDIA_ROOT) / "kits" / "tmp" / f"temp_{inquiry.id}.docx"
            temp_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(temp_path))
            template_path = temp_path
        
        # Charger le template et faire le rendu
        template = DocxTemplate(str(template_path))
        template.render(context_data)
        
        # Sauvegarder le fichier généré
        output_dir = Path(settings.PRIVATE_MEDIA_ROOT) / "drafts"
        output_dir.mkdir(parents=True, exist_ok=True)
        output_path = output_dir / f"kit_{inquiry.id}.docx"
        template.save(str(output_path))
        
        # Lire le contenu
        with open(output_path, 'rb') as f:
            docx_content = f.read()
        
        # 3) Créer/mettre à jour GeneratedDraft
        draft, created = GeneratedDraft.objects.get_or_create(
            inquiry=inquiry,
            defaults={
                "build_log": f"Généré avec succès à {timezone.now()}",
            }
        )
        
        draft.docx.save(
            f"kit_{inquiry.id}.docx",
            ContentFile(docx_content),
            save=True,
        )
        
        if not created:
            draft.build_log = f"Régénéré à {timezone.now()}"
            draft.save(update_fields=["build_log", "docx"])
        
        # Mettre à jour les statuts
        inquiry.processing_state = "DRAFT_DONE"
        inquiry.ai_status = ClientInquiry.AI_STATUS_DONE
        inquiry.ai_done_at = timezone.now()
        inquiry.save(update_fields=['processing_state', 'ai_status', 'ai_done_at'])
        
        logger.info(
            f"[build_kit_word] Draft généré avec succès pour inquiry {inquiry_id} "
            f"(taille: {len(docx_content)} bytes)"
        )
        
        # NE PAS envoyer d'email ici - le staff doit valider et publier
        
    except ClientInquiry.DoesNotExist:
        logger.error(f"Inquiry {inquiry_id} introuvable")
        raise
    except Exception as e:
        logger.exception(
            f"Erreur génération kit pour inquiry {inquiry_id}: {e}"
        )
        inquiry.processing_state = "PAID"  # Revenir à l'état précédent
        inquiry.ai_status = ClientInquiry.AI_STATUS_ERROR
        inquiry.save(update_fields=['processing_state', 'ai_status'])
        # Retry avec backoff exponentiel
        countdown = min(60 * (2 ** self.request.retries), 900)
        raise self.retry(exc=e, countdown=countdown)


def _markdown_to_docx(md_text, inquiry):
    """
    Convertit le Markdown généré en DOCX avec couverture.
    """
    from store.utils.docx_builder import build_docx_with_cover
    
    # Ajouter les informations de l'inquiry dans le Markdown
    md_with_info = f"""# Kit complet de préparation à l'audit

**Préparé pour :** {inquiry.contact_name} ({inquiry.organization_name})
**Date :** {timezone.now().strftime('%d/%m/%Y')}

---

{md_text}
"""
    
    # Générer le DOCX
    out_dir = Path(settings.PRIVATE_MEDIA_ROOT) / "kits" / "tmp"
    out_dir.mkdir(parents=True, exist_ok=True)
    
    docx_path = build_docx_with_cover(md_with_info, out_dir)
    return docx_path


def _send_kit_email(inquiry):
    """
    Envoie un email au client avec un lien signé (7 jours).
    """
    from django.core.mail import send_mail
    from django.core.signing import TimestampSigner
    from django.urls import reverse
    
    signer = TimestampSigner()
    token = signer.sign(str(inquiry.id))
    download_url = reverse('store:kit_download', args=[token])
    
    subject = (
        "Kit complet de préparation à l'audit – "
        "Votre document est prêt"
    )
    message = f"""Bonjour {inquiry.contact_name},

Votre kit complet de préparation à l'audit est prêt.

Téléchargez-le ici (lien valide 7 jours) :
{download_url}

Bien cordialement,
L'équipe AuditSansPeur
"""
    
    send_mail(
        subject,
        message,
        None,
        [inquiry.email],
        fail_silently=False
    )

