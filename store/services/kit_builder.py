# store/services/kit_builder.py
"""
Génération du "Kit de préparation à l’audit sans peur" (DOCX principalement, PDF optionnel).

Fonctions principales :
- build_docx_cover_and_guard: génère la couverture + (option) page de garde
- build_docx_content: insère Introduction, 20 Q/R, Tableau des irrégularités
- build_and_attach_kit: orchestre la génération et attache les fichiers à BonusRequest

Helpers :
- extract_text_from_pdf/docx: extraction texte du fichier uploadé (si besoin)
- convert_docx_to_pdf: conversion LibreOffice headless (facultative)
- save_in_media: sauvegarde en FileField
"""

from __future__ import annotations

import os
import re
import tempfile
import subprocess
from io import BytesIO
from pathlib import Path
from typing import List, Dict, Optional, Tuple

from django.conf import settings
from django.core.files.base import ContentFile

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

try:
    # non obligatoire (pour extraction PDF)
    import PyPDF2  # type: ignore
except Exception:
    PyPDF2 = None

try:
    # non obligatoire (pour conversion DOCX -> PDF)
    # LibreOffice doit être installé côté serveur
    HAVE_LIBREOFFICE = True
except Exception:
    HAVE_LIBREOFFICE = False


# ==============================
# === PARAMÈTRES / CONSTANTES ==
# ==============================

BRAND_NAME = "Bloom Shield Gouvernance"
BRAND_FOOTER = "© Bloom Shield Gouvernance — AuditSansPeur.com | Document personnalisé — usage interne uniquement"

DEFAULT_PRODUCT_SLUG = getattr(settings, "KITBUILDER_DEFAULT_PRODUCT_SLUG", "audit-sans-peur")
ENABLE_PDF = bool(getattr(settings, "KITBUILDER_ENABLE_PDF", False))  # activer conversion auto PDF
LIBREOFFICE_BIN = getattr(settings, "KITBUILDER_LIBREOFFICE_BIN", "soffice")  # binaire LibreOffice
MEDIA_SUBDIR = getattr(settings, "KITBUILDER_MEDIA_SUBDIR", "bonus_outputs")


# ==============================
# ======= UTILS / HELPERS ======
# ==============================

def sanitize_filename(name: str) -> str:
    name = name.strip().replace(" ", "_")
    return re.sub(r"[^A-Za-z0-9._-]+", "", name)[:128]


def ensure_media_dir(subdir: str = MEDIA_SUBDIR) -> Path:
    media_root = Path(settings.MEDIA_ROOT)
    out = media_root / subdir
    out.mkdir(parents=True, exist_ok=True)
    return out


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extraction simple depuis PDF (si PyPDF2 dispo)."""
    if PyPDF2 is None:
        return ""
    text = []
    with pdf_path.open("rb") as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            try:
                text.append(page.extract_text() or "")
            except Exception:
                continue
    return "\n".join(text).strip()


def extract_text_from_docx(docx_path: Path) -> str:
    """Extraction simple depuis DOCX."""
    try:
        d = Document(str(docx_path))
    except Exception:
        return ""
    parts = []
    for p in d.paragraphs:
        parts.append(p.text)
    # tables
    for t in d.tables:
        for row in t.rows:
            cells = [c.text for c in row.cells]
            parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def convert_docx_to_pdf(docx_path: Path, out_dir: Optional[Path] = None) -> Optional[Path]:
    """
    Convertit DOCX -> PDF via LibreOffice en mode headless.
    - Nécessite LibreOffice installé sur le serveur.
    - Retourne le chemin du PDF, ou None si échec.
    """
    if not ENABLE_PDF:
        return None

    if out_dir is None:
        out_dir = docx_path.parent

    try:
        cmd = [
            LIBREOFFICE_BIN,
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(out_dir),
            str(docx_path),
        ]
        subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        pdf_path = out_dir / (docx_path.stem + ".pdf")
        return pdf_path if pdf_path.exists() else None
    except Exception:
        return None


def set_doc_defaults(doc: Document):
    """Applique une base de styles/tailles/ marges cohérentes."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")  # compat
    style.font.size = Pt(11)

    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)

    # Footer brand
    for sec in doc.sections:
        fp = sec.footer.paragraphs[0]
        fp.text = BRAND_FOOTER
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if fp.runs:
            fp.runs[0].font.size = Pt(9)


# ==================================
# ==== GÉNÉRATION DOCX — SECTIONS ===
# ==================================

def build_docx_cover_and_guard(doc: Document,
                               client_name: str,
                               client_email: str,
                               show_guard: bool = True,
                               guard_rows: Optional[List[Tuple[str, str]]] = None):
    """
    Couverture + (option) Page de garde.
    """
    # Couverture
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Kit de préparation à l’audit sans peur")
    r.bold = True
    r.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(BRAND_NAME)
    r.font.size = Pt(16)
    r.bold = True

    doc.add_paragraph("\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Propriété de : {client_name} / {client_email}")
    r.italic = True
    r.font.size = Pt(12)

    doc.add_page_break()

    if not show_guard:
        return

    # Page de garde (facultative)
    doc.add_heading("Page de garde", level=1)
    # tableau d’infos clés
    from docx.oxml import OxmlElement  # local import
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "Champ"
    hdr[1].text = "Valeur"

    default_rows = guard_rows or [
        ("Client / Organisation", client_name),
        ("Contact principal", f"{client_name} — {client_email}"),
        ("Statut juridique", "[Ministère / Collectivité / ONG / Projet / EPIC / EPA / Autre]"),
        ("Localisation", "[Ville / Région / Pays]"),
        ("Secteur / Domaine", "[Santé / Éducation / Infrastructures / Multi-secteurs / Autre]"),
        ("Texte de référence (≤ 3 pages)", "[Titre du texte, date]"),
        ("Date de personnalisation", "[JJ/MM/AAAA]"),
        ("Confidentialité", "Usage interne — diffusion interdite sans autorisation écrite."),
    ]
    for label, value in default_rows:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value

    doc.add_paragraph("")
    p = doc.add_paragraph("Résumé exécutif (3–5 lignes)")
    p.runs[0].bold = True
    doc.add_paragraph("[Brève description du texte, objectifs, portée, points d’attention.]")
    doc.add_page_break()


def build_docx_intro(doc: Document, intro_md: str):
    doc.add_heading("Introduction", level=1)
    if not intro_md:
        intro_md = (
            "Résumé du texte fourni, définitions des notions essentielles, principes directeurs, "
            "précautions d’application et conséquences de la non-conformité."
        )
    for block in intro_md.split("\n\n"):
        doc.add_paragraph(block.strip())


def build_docx_qas(doc: Document, qas_list: List[Dict[str, str]]):
    doc.add_heading("Questionnaire de préparation (20 questions)", level=1)
    if not qas_list:
        qas_list = []
    for i, qa in enumerate(qas_list, 1):
        doc.add_heading(f"Q{i}. {qa.get('question','[Question]')}", level=2)
        if qa.get("good"):
            doc.add_paragraph(f"✅ Réponse attendue : {qa['good']}")
        if qa.get("partial"):
            doc.add_paragraph(f"⚠️ Réponse partielle : {qa['partial']}")
        if qa.get("avoid"):
            doc.add_paragraph(f"❌ À éviter : {qa['avoid']}")
        if qa.get("tip"):
            doc.add_paragraph(f"💡 Conseil : {qa['tip']}")


def build_docx_irregularities_table(doc: Document, irregularities_rows: List[Dict[str, str]]):
    doc.add_heading("Tableau des irrégularités possibles", level=1)
    headers = [
        "Irrégularité constatée",
        "Référence",
        "Acteurs concernés",
        "Dispositions pratiques",
        "Gravité",
        "Conséquences",
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    irregularities_rows = irregularities_rows or []
    for row in irregularities_rows:
        cells = table.add_row().cells
        cells[0].text = row.get("irregularity", "")
        cells[1].text = row.get("ref", "")
        cells[2].text = row.get("actors", "")
        cells[3].text = row.get("action", "")
        cells[4].text = row.get("severity", "")
        cells[5].text = row.get("impact", "")


def build_docx_content(doc: Document,
                       intro_md: str,
                       qas_list: List[Dict[str, str]],
                       irregularities_rows: List[Dict[str, str]]):
    build_docx_intro(doc, intro_md)
    build_docx_qas(doc, qas_list)
    build_docx_irregularities_table(doc, irregularities_rows)


# =========================================
# ========== ORCHESTRATION / I/O ==========
# =========================================

def assemble_docx_binary(client_name: str,
                         client_email: str,
                         intro_md: str,
                         qas_list: List[Dict[str, str]],
                         irregularities_rows: List[Dict[str, str]],
                         show_guard: bool = True) -> bytes:
    """
    Construit le DOCX en mémoire et retourne les bytes.
    """
    doc = Document()
    set_doc_defaults(doc)
    build_docx_cover_and_guard(doc, client_name, client_email, show_guard=show_guard)
    build_docx_content(doc, intro_md, qas_list, irregularities_rows)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def save_in_media(content_bytes: bytes, filename: str) -> Tuple[Path, ContentFile]:
    """
    Sauvegarde le contenu en MEDIA_ROOT / MEDIA_SUBDIR et retourne (path, ContentFile).
    """
    ensure_media_dir()
    path = ensure_media_dir() / sanitize_filename(filename)
    cf = ContentFile(content_bytes)
    return path, cf


def build_and_attach_kit(
    bonus_request,
    intro_md: str,
    qas_list: List[Dict[str, str]],
    irregularities_rows: List[Dict[str, str]],
    make_pdf: bool = ENABLE_PDF,
    show_guard: bool = True,
) -> Tuple[Optional[Path], Optional[Path]]:
    """
    Construit le DOCX (et optionnellement le PDF) puis les attache à BonusRequest.

    Params:
      - bonus_request: instance de store.models.BonusRequest
      - intro_md / qas_list / irregularities_rows : contenus prêts à insérer
      - make_pdf: True => tentative de conversion auto
    Returns:
      (docx_path, pdf_path) (peuvent être None si échec ou désactivé)
    """
    # 1) DOCX en mémoire
    client_name = bonus_request.purchaser_name or "Client"
    client_email = bonus_request.delivery_email or bonus_request.purchaser_email or ""
    data = assemble_docx_binary(
        client_name=client_name,
        client_email=client_email,
        intro_md=intro_md,
        qas_list=qas_list,
        irregularities_rows=irregularities_rows,
        show_guard=show_guard,
    )

    # 2) Sauvegarde DOCX dans FileField
    docx_filename = f"kit_{bonus_request.pk}.docx"
    docx_path, cf = save_in_media(data, docx_filename)
    bonus_request.docx_path.save(docx_path.name, cf, save=True)

    # 3) Conversion PDF optionnelle (LibreOffice)
    pdf_final_path: Optional[Path] = None
    if make_pdf:
        try:
            tmpdir = ensure_media_dir()
            # On écrit le DOCX sur disque (déjà fait par FileField, mais on re-génère si besoin)
            tmp_docx = tmpdir / docx_path.name
            tmp_docx.write_bytes(data)
            pdf_path = convert_docx_to_pdf(tmp_docx, out_dir=tmpdir)
            if pdf_path and pdf_path.exists():
                with pdf_path.open("rb") as f:
                    bonus_request.pdf_path.save(pdf_path.name, ContentFile(f.read()), save=True)
                pdf_final_path = pdf_path
                # Mets à jour statut si tu veux :
                # bonus_request.status = "READY"
                # bonus_request.save(update_fields=["status"])
        except Exception:
            # En cas d’échec conversion, on laisse simplement le DOCX
            pass

    return Path(bonus_request.docx_path.name) if bonus_request.docx_path else None, pdf_final_path


# =========================================
# =========== EXTRACTIONS (option) =========
# =========================================

def extract_uploaded_text(uploaded_file_field) -> str:
    """
    Extrait une version texte du fichier uploadé par le client (PDF/DOCX),
    utile si tu veux pré-remplir 'intro_md' ou aider à la synthèse.
    """
    if not uploaded_file_field:
        return ""

    name = uploaded_file_field.name.lower()
    path = Path(settings.MEDIA_ROOT) / uploaded_file_field.name
    if not path.exists():
        # selon le stockage, assure-toi que le fichier est bien présent
        try:
            path.parent.mkdir(parents=True, exist_ok=True)
            path.write_bytes(uploaded_file_field.read())
        except Exception:
            return ""

    if name.endswith(".pdf"):
        return extract_text_from_pdf(path)
    if name.endswith(".docx"):
        return extract_text_from_docx(path)
    return ""


# =========================================
# ========= EXEMPLE D’USAGE (views) =======
# =========================================
# from .kit_builder import build_and_attach_kit, extract_uploaded_text
#
# def generate_docx(request, pk):
#     br = get_object_or_404(BonusRequest, pk=pk)
#     # 1) Récupérer texte uploadé si besoin :
#     source_text = extract_uploaded_text(br.uploaded_text)
#     # 2) Construire intro_md / qas_list / irregularities_rows (toi ou IA)
#     intro_md = "... (résumé + notions + principes + précautions + conséquences) ..."
#     qas_list = [
#       {"question": "Votre rôle ?", "good": "Réponse attendue...", "partial": "", "avoid": "", "tip": "Soyez factuel."},
#       # ... 19 autres
#     ]
#     irregularities_rows = [
#       {"irregularity": "Paiement sans service fait", "ref": "Art. X", "actors": "Ordonnateur, Comptable", "action": "Exiger PV", "severity": "Élevé", "impact": "Perte financière"},
#       # ... ≥ 9 autres
#     ]
#     build_and_attach_kit(br, intro_md, qas_list, irregularities_rows, make_pdf=True)
#     messages.success(request, "DOCX/PDF générés.")
#     return redirect("store:bonus_admin_detail", pk=pk)
